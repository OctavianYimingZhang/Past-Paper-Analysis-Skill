"""End-to-end smoke test for the past-paper-analysis suite.

Phase D regression contract: a minimal course spec + mock mapping + mock
coverage runs through the local pipeline and produces all four output
files (md, docx, xlsx, json) with the expected structural fingerprint.

This file does not exercise LLM subagents. Those are documented in the
per-spoke SKILL.md files and tested separately via integration runs
against real fixtures. The smoke test's job is to prove the *skeleton*
still works end-to-end after the suite restructure, plus assert
structural invariants on each spoke's SKILL.md.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest
from docx import Document

# The Python implementation lives at repo root. Resolve from this file
# (skills/past-paper-orchestrator/tests/test_orchestration_smoke.py).
_REPO_ROOT = Path(__file__).resolve().parents[3]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from core.bilingual_glossary import BilingualGlossary  # noqa: E402
from core.kp_cheatsheet import build_all_cheatsheets  # noqa: E402
from core.pattern_coverage import PatternCoverage, PatternOccurrence  # noqa: E402
from core.sensitivity import SensitivityCell, SensitivitySweep  # noqa: E402
from core.statistical_model import KPPosterior  # noqa: E402
from scripts.report_writer import write_docx, write_excel, write_json, write_markdown  # noqa: E402


# ---------------------------------------------------------------------------
# Tiny fixture factories (kept minimal; richer versions live in the existing
# tests/test_report_writer_docx.py).
# ---------------------------------------------------------------------------


def _posterior(
    kp_id: str,
    tier: str,
    posterior_mean: float,
    n_papers: int = 5,
    raw_hits: int = 5,
) -> KPPosterior:
    return KPPosterior(
        kp_id=kp_id,
        n_papers=n_papers,
        raw_hits=raw_hits,
        weighted_hits=float(raw_hits),
        weighted_N=float(n_papers),
        lambda_used=0.2,
        tau_used=1.0,
        reference_year=2026,
        coverage_share=0.1,
        prior_alpha=0.1,
        prior_beta=0.9,
        posterior_alpha=raw_hits + 0.1,
        posterior_beta=n_papers - raw_hits + 0.9,
        posterior_mean=posterior_mean,
        ci_lower_95=max(0.0, posterior_mean - 0.3),
        ci_upper_95=min(1.0, posterior_mean + 0.15),
        hotness_mean_share=0.1,
        hotness_std_share=0.02,
        trend_label="stable",
        trend_delta=0.0,
        trend_ci_95=(-0.1, 0.1),
        historical_mean=raw_hits / n_papers if n_papers else 0.0,
        tier=tier,
        tier_reasons=(f"posterior_mean={posterior_mean:.2f}",),
        sensitivity_band="stable",
        warnings=(),
    )


def _coverage(kp_id: str) -> PatternCoverage:
    return PatternCoverage(
        kp_id=kp_id,
        pattern_id=f"{kp_id}.P01",
        raw_hits=3,
        weighted_hits=2.5,
        last_seen_year=2024.4,
        first_seen_year=2020.4,
        inter_arrival_years_mean=1.0,
        inter_arrival_years_max=1.0,
        saturation_index=0.4,
        freshness_flag=False,
        predicted_score=0.8,
        complications_seen=(),
        complications_unseen=(),
        occurrences=(
            PatternOccurrence(
                year=2024.4,
                question_number="1",
                confidence=0.9,
                is_primary=True,
                complications=(),
            ),
        ),
        warnings=(),
        tier="hot",
        tier_reasons=("hot",),
    )


def _sweep(kp_id: str) -> SensitivitySweep:
    cell = SensitivityCell(
        lam=0.2,
        tau=1.0,
        posterior_mean=0.85,
        ci_lower_95=0.55,
        ci_upper_95=0.99,
        tier="anchor",
        warnings=(),
    )
    return SensitivitySweep(
        kp_id=kp_id,
        cells=(cell,),
        distinct_tiers=("anchor",),
        band="stable",
    )


# ---------------------------------------------------------------------------
# End-to-end smoke
# ---------------------------------------------------------------------------


def test_skeleton_renders_all_four_files(tmp_path: Path) -> None:
    """Orchestrator skeleton produces md / docx / xlsx / json end-to-end."""
    posteriors = [
        _posterior("L01.01", tier="anchor", posterior_mean=0.85, raw_hits=5),
        _posterior("L02.01", tier="core", posterior_mean=0.42, raw_hits=2),
        _posterior("L03.01", tier="oneoff", posterior_mean=0.18, raw_hits=1),
    ]
    pattern_coverage = [_coverage("L01.01")]
    sweeps = {p.kp_id: _sweep(p.kp_id) for p in posteriors}
    loo: dict = {}

    course_meta = {
        "course_id": "skeleton-smoke",
        "course_name": "Skeleton Smoke Test",
        "reference_year": 2026,
        "n_papers": 5,
        "n_kps": len(posteriors),
    }
    hyperparameters = {
        "lambda": 0.2,
        "tau": 1.0,
        "alpha": 0.3,
        "reference_year": 2026,
    }

    md_path = tmp_path / "skeleton-smoke-analysis.md"
    docx_path = tmp_path / "skeleton-smoke-analysis.docx"
    xlsx_path = tmp_path / "skeleton-smoke-analysis.xlsx"
    json_path = tmp_path / "skeleton-smoke-analysis.json"

    write_markdown(
        md_path,
        posteriors=posteriors,
        sweeps=sweeps,
        hyperparameters=hyperparameters,
        pattern_coverage=pattern_coverage,
        course_meta=course_meta,
        loo=loo,
    )
    write_docx(
        docx_path,
        posteriors=posteriors,
        sweeps=sweeps,
        hyperparameters=hyperparameters,
        pattern_coverage=pattern_coverage,
        course_meta=course_meta,
        loo=loo,
    )
    write_excel(
        xlsx_path,
        posteriors=posteriors,
        sweeps=sweeps,
        loo=loo,
        hyperparameters=hyperparameters,
    )
    write_json(
        json_path,
        posteriors=posteriors,
        sweeps=sweeps,
        loo=loo,
        hyperparameters=hyperparameters,
    )

    for path in (md_path, docx_path, xlsx_path, json_path):
        assert path.exists(), f"missing output: {path}"
        assert path.stat().st_size > 0, f"empty output: {path}"

    doc = Document(str(docx_path))
    assert len(doc.paragraphs) > 0
    assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)

    payload = json.loads(json_path.read_text(encoding="utf-8"))
    assert "posteriors" in payload
    assert {p["kp_id"] for p in payload["posteriors"]} == {"L01.01", "L02.01", "L03.01"}


def test_kp_cheatsheet_builder_runs_against_smoke_inputs() -> None:
    """The cheat-sheet builder consumes the same shape the orchestrator hands cheatsheet-writer."""
    posteriors = [
        _posterior("L01.01", tier="anchor", posterior_mean=0.85, raw_hits=5),
        _posterior("L02.01", tier="core", posterior_mean=0.42, raw_hits=2),
        _posterior("L03.01", tier="oneoff", posterior_mean=0.18, raw_hits=1),
    ]
    sheets = build_all_cheatsheets(
        posteriors=posteriors,
        pattern_coverages=[_coverage("L01.01")],
        pattern_definitions=[],
        mapping_questions=[],
        kps=[
            {"id": "L01.01", "label": "Anchor KP"},
            {"id": "L02.01", "label": "Core KP"},
            {"id": "L03.01", "label": "Oneoff KP"},
        ],
        narratives={},
    )
    assert set(sheets) == {"L01.01", "L02.01", "L03.01"}
    assert sheets["L01.01"].tier == "anchor"


def test_bilingual_glossary_round_trip(tmp_path: Path) -> None:
    """Glossary persists across runs — used by the bilingual rendering layer."""
    path = tmp_path / "glossary.json"
    g1 = BilingualGlossary(path=path)
    g1.register("posterior mean", {"zh": "后验均值"}, source="methodology.md")
    g1.dump()

    g2 = BilingualGlossary.load(path)
    assert g2.lookup("posterior mean", "zh") == "后验均值"


# ---------------------------------------------------------------------------
# Suite-shape invariants
# ---------------------------------------------------------------------------


def test_orchestrator_skill_md_exists_and_lists_seven_spokes() -> None:
    """Orchestrator SKILL.md mentions every spoke + parallel dispatch."""
    skill_md = _REPO_ROOT / "skills" / "past-paper-orchestrator" / "SKILL.md"
    assert skill_md.exists()
    text = skill_md.read_text(encoding="utf-8")
    for spoke in (
        "paper-ingest",
        "kp-pattern-mapper",
        "stat-engine",
        "cheatsheet-writer",
        "drill-curator",
        "technique-coach",
        "report-renderer",
    ):
        assert spoke in text, f"orchestrator missing spoke reference: {spoke}"
    assert "PARALLEL" in text or "parallel" in text


@pytest.mark.parametrize(
    "spoke",
    [
        "paper-ingest",
        "kp-pattern-mapper",
        "stat-engine",
        "cheatsheet-writer",
        "drill-curator",
        "technique-coach",
        "report-renderer",
    ],
)
def test_each_spoke_has_skill_md_with_frontmatter(spoke: str) -> None:
    """Every spoke SKILL.md has YAML frontmatter (name + description + triggers)."""
    skill_md = _REPO_ROOT / "skills" / spoke / "SKILL.md"
    assert skill_md.exists(), f"missing SKILL.md: {skill_md}"
    text = skill_md.read_text(encoding="utf-8")
    assert text.startswith("---\n"), f"{spoke}: SKILL.md must start with frontmatter"
    assert f"name: {spoke}" in text, f"{spoke}: SKILL.md frontmatter must declare name"
    assert "description:" in text
    assert "triggers:" in text


def test_orchestrator_references_present() -> None:
    """The Phase C-introduced references exist."""
    refs_dir = _REPO_ROOT / "skills" / "past-paper-orchestrator" / "references"
    for filename in (
        "voice-and-conviction.md",
        "report-format.md",
        "companion-skills.md",
        "methodology.md",
        "tier-definitions.md",
        "course-spec-schema.md",
        "presets.md",
        "subagent-orchestration.md",
    ):
        assert (refs_dir / filename).exists(), f"missing reference: {filename}"


def test_external_borrowings_log_has_required_columns() -> None:
    """references/external-borrowings.md must list license + verification per row."""
    log = _REPO_ROOT / "references" / "external-borrowings.md"
    assert log.exists()
    text = log.read_text(encoding="utf-8")
    # Spot-check the header columns.
    for column in ("Source repo", "License", "Verified on", "Carrier file"):
        assert column in text, f"borrowings log missing column: {column}"


def test_new_agent_prompts_exist() -> None:
    """The two new Phase B agent prompts (drill-curator + technique-coach) exist."""
    drill_prompt = _REPO_ROOT / "skills" / "drill-curator" / "agents" / "drill-curator.md"
    technique_prompt = _REPO_ROOT / "skills" / "technique-coach" / "agents" / "technique-coach.md"
    assert drill_prompt.exists()
    assert technique_prompt.exists()
    for prompt in (drill_prompt, technique_prompt):
        text = prompt.read_text(encoding="utf-8")
        assert text.startswith("---\n"), f"{prompt.name}: must start with frontmatter"
        assert "model:" in text and "opus-4.7" in text.lower()
